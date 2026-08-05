from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "informe" / "informe.md"
OUTPUT = ROOT / "informe" / "Felizia.TFI.SCI2025.docx"
REFERENCE = ROOT / "output" / "informe_reference.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 98, 108)
LIGHT_FILL = "F2F4F7"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, *, size, color, before, after, bold=True, line=1.1):
    style.font.name = "Calibri"
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line
    style.paragraph_format.keep_with_next = True


def add_field(run, instruction, fallback):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def create_reference_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.widow_control = True

    configure_style(doc.styles["Title"], size=25, color=INK, before=0, after=9, line=1.0)
    configure_style(doc.styles["Subtitle"], size=15, color=DARK_BLUE, before=0, after=18, bold=False, line=1.05)
    configure_style(doc.styles["Heading 1"], size=16, color=BLUE, before=16, after=8)
    configure_style(doc.styles["Heading 2"], size=13, color=BLUE, before=12, after=6)
    configure_style(doc.styles["Heading 3"], size=12, color=DARK_BLUE, before=8, after=4)

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9.5)
    caption.font.color.rgb = MUTED
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("Sistemas de Control I  |  TFI - Horno de cataforesis"), size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_font(footer.add_run("Página "), size=9, color=MUTED)
    page_run = footer.add_run()
    set_font(page_run, size=9, color=MUTED)
    add_field(page_run, "PAGE", "1")

    doc.add_paragraph("Reference document.")
    REFERENCE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REFERENCE)


def table_widths(headers):
    key = tuple(headers)
    exact = {
        ("Símbolo", "Descripción", "Unidad"): [1800, 5500, 2060],
        ("Señal", "Rango de ingeniería", "Nivel físico adoptado"): [2500, 2500, 4360],
        ("Parámetro", "Valor", "Procedencia"): [2300, 1800, 5260],
        ("Ensayo", "Especificación", "Criterio"): [2700, 4300, 2360],
        ("Magnitud", "Especificación", "Resultado", "Cumple"): [3000, 2100, 2100, 2160],
        ("Archivo", "Función principal"): [2700, 6660],
    }
    if key in exact:
        return exact[key]
    widths = [TABLE_WIDTH // len(headers)] * len(headers)
    widths[-1] += TABLE_WIDTH - sum(widths)
    return widths


def patch_document(path):
    doc = Document(path)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

    cover_lines = {
        "UNIVERSIDAD NACIONAL DE CÓRDOBA",
        "FACULTAD DE CIENCIAS EXACTAS, FÍSICAS Y NATURALES",
        "Carrera: Ingeniería Electrónica",
        "Cátedra: Sistemas de Control I",
        "Docentes: Ing. Adrián Claudio Agüero e Ing. Juan Pablo Pedroni",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in cover_lines:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.05
            for run in paragraph.runs:
                set_font(run, size=11, color=INK, bold=text.startswith(("UNIVERSIDAD", "FACULTAD")))
        if text == "UNIVERSIDAD NACIONAL DE CÓRDOBA":
            paragraph.paragraph_format.space_before = Pt(70)
        if paragraph._p.xpath(".//m:oMathPara"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(7)
        if paragraph.style.name == "Caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = False
            paragraph.paragraph_format.keep_together = True
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(0)
        if paragraph._p.xpath("./w:pPr/w:numPr"):
            paragraph.paragraph_format.keep_with_next = False
            paragraph.paragraph_format.keep_together = False
            paragraph.paragraph_format.widow_control = True
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True

    for table in doc.tables:
        table.style = "Table Grid"
        table.autofit = False
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        widths = table_widths(headers)
        tbl_pr = table._tbl.tblPr

        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
        tbl_w.set(qn("w:type"), "dxa")

        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
        tbl_ind.set(qn("w:type"), "dxa")

        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(width))
            grid.append(grid_col)

        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            for existing_header in list(tr_pr.findall(qn("w:tblHeader"))):
                tr_pr.remove(existing_header)
            if row_index == 0:
                header_flag = OxmlElement("w:tblHeader")
                header_flag.set(qn("w:val"), "true")
                tr_pr.append(header_flag)
            for col_index, cell in enumerate(row.cells):
                set_cell_width(cell, widths[col_index])
                set_cell_margins(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if row_index == 0:
                    shade_cell(cell, LIGHT_FILL)
                for p in cell.paragraphs:
                    # Pandoc emits a paragraph style named Compact for table
                    # cells, but that style is not present in Word reference
                    # documents. LibreOffice then lays out cell contents as if
                    # they were outside the table. Normal is explicit and
                    # portable across both renderers.
                    p.style = doc.styles["Normal"]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.05
                    if col_index > 0 and len(headers) >= 3:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        set_font(run, size=9.2, bold=True if row_index == 0 else None)

    # Ask Word/LibreOffice to refresh TOC and page-number fields on open/export.
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.core_properties.title = "Control de temperatura de una zona equivalente de un horno de cataforesis"
    doc.core_properties.subject = "Trabajo Final Integrador - Sistemas de Control I"
    doc.core_properties.author = "Valentín Felizia"
    doc.core_properties.keywords = "control, horno, cataforesis, PI, Octave"
    doc.save(path)


def build():
    missing_figures = [
        name for name in (
            "00_diagrama_bloques.png",
            "01_step_planta_Gp.png",
            "02_polos_planta.png",
            "03_step_perturbacion_GL.png",
            "04_Gp_vs_Gred.png",
            "05_lugar_raices_PI.png",
            "06_referencia_P_vs_PI.png",
            "07_perturbacion_PI.png",
            "08_bode_lazo_abierto_PI.png",
        ) if not (ROOT / "figuras" / name).exists()
    ]
    if missing_figures:
        raise RuntimeError(
            "Faltan figuras. Ejecutar octave/run_all.m y generar 00_diagrama_bloques.png: "
            + ", ".join(missing_figures)
        )

    create_reference_docx()
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        raise RuntimeError("pandoc no está disponible en PATH")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="tfi_informe_") as temp_dir:
        draft = Path(temp_dir) / "draft.docx"
        subprocess.run(
            [
                pandoc,
                str(SOURCE),
                "--from=markdown+tex_math_dollars+raw_attribute",
                "--to=docx",
                "--number-sections",
                f"--reference-doc={REFERENCE}",
                "--resource-path",
                str(ROOT / "informe") + os.pathsep + str(ROOT),
                "--output",
                str(draft),
            ],
            check=True,
            cwd=ROOT,
        )
        try:
            shutil.copy2(draft, OUTPUT)
        except PermissionError:
            alt = OUTPUT.with_name(OUTPUT.stem + "_nuevo.docx")
            shutil.copy2(draft, alt)
            raise PermissionError(
                f"No se pudo sobrescribir {OUTPUT.name} (archivo abierto). "
                f"Se generó {alt.name}; cerrá Word y renombralo."
            ) from None

    patch_document(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
