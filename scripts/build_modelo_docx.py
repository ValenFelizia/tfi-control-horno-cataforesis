from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "modelo_matematico.md"
REFERENCE = ROOT / "output" / "modelo_reference.docx"
OUTPUT = ROOT / "docs" / "modelo_matematico.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, *, size, color, before, after, bold=True, line=1.1):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line
    style.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def create_reference_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    configure_style(
        doc.styles["Title"],
        size=24,
        color=RGBColor(11, 37, 69),
        before=0,
        after=5,
        bold=True,
        line=1.0,
    )
    configure_style(
        doc.styles["Subtitle"],
        size=14,
        color=DARK_BLUE,
        before=0,
        after=10,
        bold=False,
        line=1.05,
    )
    configure_style(
        doc.styles["Heading 1"],
        size=16,
        color=BLUE,
        before=16,
        after=8,
    )
    configure_style(
        doc.styles["Heading 2"],
        size=13,
        color=BLUE,
        before=12,
        after=6,
    )
    configure_style(
        doc.styles["Heading 3"],
        size=12,
        color=DARK_BLUE,
        before=8,
        after=4,
    )

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("Sistemas de Control I  |  TFI - Modelado")
    set_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    add_page_field(footer)

    # Pandoc imports document-level styles and section properties from here.
    doc.add_paragraph("Reference document.")
    REFERENCE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REFERENCE)


def make_pandoc_source() -> str:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith(">"))
    body = "\n".join(lines[start:])
    metadata = """---
title: "Cuaderno de modelado matemático"
subtitle: "Control de temperatura de una zona equivalente de un horno de cataforesis"
author: "Valentín Felizia"
date: "Modelo nominal y PI | Cursado 2025 | Revisión 3 de agosto de 2026"
lang: es-AR
---

"""
    return metadata + body + "\n"


def patch_document():
    doc = Document(OUTPUT)
    source_section = False

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "13. Fuentes académicas usadas en esta etapa":
            source_section = True
            paragraph.paragraph_format.page_break_before = True
        if paragraph._p.xpath(".//m:oMathPara"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(7)
        if paragraph.style.name in {"Block Text", "Intense Quote"}:
            p_pr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), CALLOUT_FILL)
            p_pr.append(shd)
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.18)
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(10)
        if paragraph.text.strip() in {"1.1 Variables", "5.1 Valores provisionales"}:
            paragraph.paragraph_format.keep_with_next = False
        if paragraph._p.xpath("./w:pPr/w:numPr"):
            paragraph.style = doc.styles["Normal"]
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.keep_with_next = False
            paragraph.paragraph_format.keep_together = False
            paragraph.paragraph_format.widow_control = False
        if source_section and paragraph.style.name == "Normal":
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_font(run, size=9.5)

    for table in doc.tables:
        table.style = "Table Grid"
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
        tbl_w.set(qn("w:type"), "dxa")

        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
        tbl_ind.set(qn("w:type"), "dxa")

        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers == ["Símbolo", "Descripción", "Unidad"]:
            widths = [1800, 5500, 2060]
        elif headers == ["Parámetro", "Valor", "Procedencia"]:
            widths = [2500, 1800, 5060]
        else:
            widths = [TABLE_WIDTH // len(headers)] * len(headers)
            widths[-1] += TABLE_WIDTH - sum(widths)

        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(width))
            grid.append(grid_col)

        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            header_flag = tr_pr.find(qn("w:tblHeader"))
            if header_flag is not None:
                tr_pr.remove(header_flag)
            if row_index == 0:
                header_flag = OxmlElement("w:tblHeader")
                header_flag.set(qn("w:val"), "true")
                tr_pr.append(header_flag)
            for col_index, cell in enumerate(row.cells):
                set_cell_width(cell, widths[col_index])
                set_cell_margins(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if row_index == 0:
                    shade_cell(cell, LIGHT_FILL)
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles["Normal"]
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.05
                    if col_index == len(widths) - 1 and row_index > 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_font(
                            run,
                            size=9.5,
                            bold=True if row_index == 0 else None,
                        )

    doc.core_properties.title = "Cuaderno de modelado matemático"
    doc.core_properties.subject = "TFI de Sistemas de Control I"
    doc.core_properties.author = "Valentín Felizia"
    doc.core_properties.keywords = "control, horno, cataforesis, Octave"
    doc.save(OUTPUT)


def build():
    create_reference_docx()
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        raise RuntimeError("pandoc no está disponible en PATH")
    with tempfile.TemporaryDirectory(prefix="tfi_docx_") as temp_dir:
        temp_markdown = Path(temp_dir) / "modelo.md"
        temp_markdown.write_text(make_pandoc_source(), encoding="utf-8")
        subprocess.run(
            [
                pandoc,
                str(temp_markdown),
                "--from=markdown+tex_math_dollars+tex_math_single_backslash",
                "--to=docx",
                f"--reference-doc={REFERENCE}",
                "--output",
                str(OUTPUT),
            ],
            check=True,
        )
    patch_document()


if __name__ == "__main__":
    build()
